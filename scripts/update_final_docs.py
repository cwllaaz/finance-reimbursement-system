from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
TODAY = datetime.now().strftime("%Y-%m-%d")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True)
        set_cell_shading(header_cells[index], "F2F4F7")
        if widths:
            header_cells[index].width = Inches(widths[index])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            if widths:
                cells[index].width = Inches(widths[index])
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def setup_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.name = "Calibri"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor.from_string("555555")

    doc.add_paragraph()
    return doc


roles = [
    ("EMPLOYEE", "员工", "提交报销、申购、劳务、借款等申请，查看自己的申请。"),
    ("DEPARTMENT_MANAGER", "部门负责人", "查看本部门数据，处理本部门待审批事项。"),
    ("FINANCE", "财务人员", "财务初审、复核、预算、收入登记、总台账和导出。"),
    ("OFFICE", "办公室", "申购相关协作，资产验收入库和资产台账维护。"),
    ("EXECUTIVE", "执行院长", "查看全院财务数据，处理院长审批节点。"),
    ("CASHIER", "出纳", "处理付款任务，登记付款日期、金额和凭证号。"),
    ("ADMIN", "管理员", "用户管理、系统维护、全量数据查看和安全审计。"),
]

modules = [
    ("首页仪表盘", "财务/领导可见，员工不可见；部门负责人只看本部门。"),
    ("我的申请", "由后端统一工作台返回本人创建的报销、申购、劳务和借款。"),
    ("我的待办", "按当前角色、审批节点和部门范围返回真实待处理任务。"),
    ("已办事项", "根据审批记录返回当前用户实际审批、付款或复核过的事项。"),
    ("报销管理", "审批编号 BX+日期+序号；OCR 金额校验；5 万以上会议材料校验。"),
    ("申购管理", "编号 CG+日期+序号；多明细；5 万以上院务会材料校验。"),
    ("资产出入库/验收", "编号 LY/ZC；办公室入库，使用人领用，保管人和位置可追踪。"),
    ("劳务/酬金发放", "编号 LW+日期+序号；多人领款；金额大写；付款与复核。"),
    ("暂借款/预付款", "编号 YF+日期+序号；待还款、部分冲账、已冲账、逾期提醒。"),
    ("收入登记", "记录收款日期、凭证号、缴款方、收入分类、到账账户和发票附件。"),
    ("财务总台账", "汇总收入、报销、劳务、申购、借款支出，支持筛选和 Excel 导出。"),
    ("操作日志/安全审计", "记录登录、审批、导出、用户管理、附件上传等关键操作。"),
    ("登录与附件安全", "BCrypt 密码、限时 token、密码修改后会话失效、受权限保护的附件下载。"),
]

accounts = [
    ("employee", "123456", "员工", "提交和查看个人申请"),
    ("manager", "123456", "部门负责人", "处理本部门审批"),
    ("finance", "123456", "财务人员", "初审、复核、预算、收入、台账"),
    ("office", "123456", "办公室", "申购协作、资产验收"),
    ("executive", "123456", "执行院长", "全院数据和院长审批"),
    ("cashier", "123456", "出纳", "付款任务"),
    ("admin", "123456", "管理员", "全量管理和审计"),
]


def feature_summary():
    doc = setup_doc("内部凭证财务系统功能汇总", f"最终整合验收版 | {TODAY}")
    doc.add_heading("一、系统定位", level=1)
    doc.add_paragraph("本系统以研究院内部凭证流转为核心，覆盖申请、审批、付款、复核、预算扣减、资产台账、收入登记、收支总台账和安全审计，适合作为课程项目和演示型业务系统。")

    doc.add_heading("二、功能模块", level=1)
    add_table(doc, ["模块", "功能说明"], modules, [1.8, 4.5])

    doc.add_heading("三、角色权限", level=1)
    add_table(doc, ["角色代码", "角色名称", "权限说明"], roles, [1.5, 1.3, 3.6])

    doc.add_heading("四、核心业务规则", level=1)
    add_bullets(doc, [
        "报销审批流程：申请人提交 -> 财务初审 -> 部门负责人审批 -> 执行院长审批 -> 出纳付款并上传银行回执 -> 财务复核 -> 完成。",
        "报销单自动生成 BX 编号；申购单生成 CG 编号；劳务单生成 LW 编号；借款/预付款生成 YF 编号；资产生成 LY/ZC 编号；收入生成 SR 编号。",
        "金额超过 5 万元的报销和申购必须填写说明并上传会议审议材料，否则不能继续提交或审批。",
        "OCR 识别结果需要上传人确认，系统会比对发票金额与报销金额，异常时给出红色提醒，但不强制阻断审批。",
        "财务复核通过后才扣减预算，避免审批未完成时提前影响预算数据。",
        "所有关键操作写入操作日志，管理员可按模块、操作和关键词查询。",
        "旧明文密码在系统启动时自动升级为 BCrypt；附件限制为 10MB 并校验扩展名、MIME 和危险路径。",
    ])

    doc.add_heading("五、验收结果", level=1)
    add_bullets(doc, [
        "前端构建通过，菜单已整合为仪表盘、工作台、业务模块、台账和系统管理。",
        "后端 51 项测试全部通过并完成打包，覆盖业务流程、角色权限、安全机制和演示数据幂等性。",
        "演示账号覆盖员工、部门负责人、财务、办公室、执行院长、出纳、管理员。",
        "报销、申购、劳务、借款均准备草稿、审批中、待后续处理和已完成演示数据；完成单含真实审批记录。",
    ])
    doc.save(ROOT / "财务报销与预算管理系统功能汇总.docx")


def user_manual():
    doc = setup_doc("内部凭证财务系统使用文档", f"最终整合验收版 | {TODAY}")
    doc.add_heading("一、启动系统", level=1)
    add_numbered(doc, [
        "打开后端：进入 backend 目录，运行 mvnw.cmd spring-boot:run，确认 8080 端口启动成功。",
        "打开前端：进入 frontend 目录，运行 npm run dev，浏览器访问前端地址。",
        "如果是服务器部署版，直接访问 http://服务器公网IP/finance。",
    ])

    doc.add_heading("二、演示账号", level=1)
    add_table(doc, ["用户名", "密码", "角色", "演示用途"], accounts, [1.3, 1.0, 1.2, 3.0])

    doc.add_heading("三、常用演示流程", level=1)
    add_numbered(doc, [
        "使用 employee 登录，进入我的申请或报销管理，新建报销单并上传发票附件。",
        "打开报销详情，执行 OCR 识别，检查金额校验提示，确认 OCR 数据。",
        "提交报销单，切换 finance、manager、executive、cashier 按流程审批和付款。",
        "finance 完成财务复核后，查看预算扣减和审批时间轴。",
        "使用 admin 登录查看操作日志，确认关键操作已经被记录。",
        "使用 finance 登录收入登记，新增收入记录并上传附件。",
        "进入财务总台账，按日期、部门、业务类型筛选，并导出 Excel。",
        "可直接使用固定演示编号 BX/CG/LW/YF20260101001-004 查看不同阶段和完整时间轴。",
    ])

    doc.add_heading("四、页面说明", level=1)
    add_table(doc, ["页面", "使用说明"], modules, [1.8, 4.5])

    doc.add_heading("五、注意事项", level=1)
    add_bullets(doc, [
        "员工不能查看财务仪表盘和总台账，这是前后端同时控制的权限。",
        "部门负责人只能查看本部门数据。",
        "5 万元以上报销/申购必须上传会议审议材料。",
        "出纳付款前需要上传银行回执，并填写付款日期、金额、凭证号。",
        "附件必须通过系统内“查看”按钮下载，不能直接访问 /uploads 路径。",
        "登录 token 默认 120 分钟有效；退出登录或修改密码后旧 token 立即失效。",
        "部署到服务器后，不建议开放 MySQL 3306 公网端口。",
    ])
    doc.save(ROOT / "财务项目使用文档.docx")


def project_report():
    doc = setup_doc("财务项目汇报", f"Vibe Coding 项目最终汇报 | {TODAY}")
    doc.add_heading("一、项目背景", level=1)
    doc.add_paragraph("本项目从课程第一次作业出发，最初目标是完成一个财务报销与预算管理系统，后续根据老师意见扩展为更贴近研究院内部凭证管理的综合系统。")

    doc.add_heading("二、我如何使用 Vibe Coding", level=1)
    add_bullets(doc, [
        "先把老师要求拆成可执行任务，例如角色权限、报销细化、审批流程、申购、资产、劳务、借款、收入和总台账。",
        "每次只给 AI 一个明确任务，要求同时修改数据库、后端、前端、权限、日志并运行构建验证。",
        "遇到报错时把截图或终端输出发给 AI，由 AI 帮我定位原因并给出下一步命令。",
        "我负责确认业务是否符合老师意见、测试页面流程、配置服务器和 API Key；AI 负责大部分代码实现和文档整理。",
        "最终通过多轮迭代，把一个简单作业扩展成完整演示系统。",
    ])

    doc.add_heading("三、系统成果", level=1)
    add_table(doc, ["类别", "完成内容"], [
        ("业务范围", "报销、申购、资产、劳务、借款、收入、预算、台账、日志。"),
        ("流程能力", "财务初审、部门审批、执行院长审批、出纳付款、财务复核。"),
        ("展示能力", "首页仪表盘、工作台、时间轴、图表、Excel 导出、OCR 校验。"),
        ("权限能力", "7 类角色，菜单权限和后端 API 权限同步控制。"),
        ("安全能力", "BCrypt 密码迁移、限时 token、会话失效、文件类型校验和授权下载。"),
        ("演示能力", "四类业务覆盖草稿、审批中、待付款/待后续处理、已完成，并配套审批记录和附件。"),
        ("部署能力", "支持本地运行，也支持服务器 Nginx + Spring Boot + MySQL 部署。"),
    ], [1.4, 4.9])

    doc.add_heading("四、验收说明", level=1)
    add_bullets(doc, [
        "后端 51 项测试全部通过并完成打包验证。",
        "前端已运行生产构建验证。",
        "系统保留演示账号和演示数据，方便老师登录体验。",
        "系统已完成 BCrypt 密码、限时 token 和安全附件下载；若作为真实商用系统，还需补充消息通知、自动备份、灾备和更完整的合规审计。",
    ])
    doc.save(ROOT / "财务项目汇报.docx")


def web_manual():
    doc = setup_doc("财务系统 Web 端访问使用文档", f"服务器访问版 | {TODAY}")
    doc.add_heading("一、访问地址", level=1)
    doc.add_paragraph("部署完成后，浏览器访问：http://服务器公网IP/finance。后端健康检查地址为：http://服务器公网IP/api/health。")

    doc.add_heading("二、服务器组件", level=1)
    add_table(doc, ["组件", "作用"], [
        ("Nginx", "对外提供 80/443 访问，转发 /api 到后端，托管 /finance 前端静态文件。"),
        ("Spring Boot", "运行后端 API、业务逻辑、权限校验和文件上传。"),
        ("MySQL", "保存用户、部门、报销、审批、附件、资产、收入和日志数据。"),
        ("系统服务", "使用 systemd 管理后端进程，保证重启后可恢复。"),
    ], [1.5, 4.8])

    doc.add_heading("三、更新部署流程", level=1)
    add_numbered(doc, [
        "本地运行前端构建 npm run build。",
        "本地运行后端打包 mvnw.cmd clean test package。",
        "上传 backend/target/backend-0.0.1-SNAPSHOT.jar、frontend/dist、deploy 配置和 database/init.sql 到服务器。",
        "服务器执行数据库初始化或迁移，重启后端服务。",
        "复制前端 dist 到 Nginx /finance 目录，重载 Nginx。",
        "访问 /finance、/api/health 验证前后端正常。",
    ])

    doc.add_heading("四、上线安全建议", level=1)
    add_bullets(doc, [
        "服务器只开放 22、80、443，MySQL 不建议公网开放。",
        "使用强密码或 SSH 密钥登录服务器。",
        "定期备份数据库和上传附件目录。",
        "系统已启用 BCrypt 密码和授权附件下载；真实上线还应启用 HTTPS、接口限流和日志异地备份。",
    ])
    doc.save(ROOT / "财务系统Web端访问使用文档.docx")


if __name__ == "__main__":
    feature_summary()
    user_manual()
    project_report()
    web_manual()
    print("updated final documents")
