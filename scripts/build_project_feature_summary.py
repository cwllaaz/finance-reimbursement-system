from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = r"C:\Users\13289\Desktop\极客工坊文件\第一次作业\finance-reimbursement-system\财务报销与预算管理系统功能汇总.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D9E2F3"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(46, 116, 181)
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(31, 77, 120)
    return paragraph


def add_body_paragraph(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        rest = text[len(bold_prefix):]
        paragraph.add_run(rest)
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(11)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)


def add_feature_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Cm(3.2), Cm(5.4), Cm(7.7)]
    headers = ["功能模块", "已实现功能", "展示价值"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = widths[index]
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, header, bold=True, color="1F4D78")

    for module, features, value in rows:
        cells = table.add_row().cells
        values = [module, features, value]
        for index, text in enumerate(values):
            cells[index].width = widths[index]
            set_cell_text(cells[index], text)
    set_table_borders(table)
    doc.add_paragraph()
    return table


def add_roles_table(doc):
    rows = [
        ("employee", "员工", "提交、编辑、查看自己的报销单；上传发票；确认 OCR 信息。"),
        ("manager", "部门负责人", "查看本部门待审批报销单；进行部门审批；查看审批流程。"),
        ("finance", "财务人员", "财务审批、预算管理、数据统计、Excel 导出。"),
        ("admin", "管理员", "查看全部数据；管理用户；查看操作日志；维护系统数据。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    headers = ["账号", "角色", "主要权限"]
    widths = [Cm(3), Cm(3), Cm(10.3)]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_shading(cell, "F2F4F7")
        set_cell_text(cell, header, bold=True, color="1F4D78")
    for account, role, permission in rows:
        cells = table.add_row().cells
        for i, text in enumerate([account, role, permission]):
            cells[i].width = widths[i]
            set_cell_text(cells[i], text)
    set_table_borders(table, "DADCE0")


def main():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("财务报销与预算管理系统功能汇总")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("Vue3 + Spring Boot + MySQL + 百度 OCR API | Vibe Coding 实践项目")
    subtitle_run.font.name = "Microsoft YaHei"
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(102, 112, 133)

    add_heading(doc, "一、项目概述", 1)
    add_body_paragraph(
        doc,
        "本项目是一个面向企业内部费用报销场景的后台管理系统，围绕“报销申请、审批流转、预算控制、发票识别、数据统计、安全审计”形成完整闭环。系统支持员工、部门负责人、财务人员、管理员四类角色，既能完成基础 CRUD，也能展示较完整的业务流程和数据安全意识。"
    )
    add_body_paragraph(
        doc,
        "项目定位：大学生课程实践项目，重点展示通过 Vibe Coding 与 AI 辅助，从需求拆解、环境搭建、前后端开发、数据库设计、上线部署到功能迭代的完整实现过程。",
        bold_prefix="项目定位：",
    )

    add_heading(doc, "二、技术架构", 1)
    add_feature_table(doc, [
        ("前端", "Vue3、Element Plus、ECharts、Axios", "形成后台管理系统界面，支持表格、弹窗、图表、时间轴、文件上传等交互。"),
        ("后端", "Spring Boot、Spring Data JPA、REST API", "负责业务规则、权限控制、审批流、OCR 校验、Excel 导出等核心逻辑。"),
        ("数据库", "MySQL", "存储用户、部门、预算、报销单、审批记录、附件、OCR 结果、操作日志等数据。"),
        ("第三方能力", "百度 OCR API", "识别发票号码、发票代码、开票日期、金额、税额、销售方、购买方等字段。"),
        ("部署", "Ubuntu 服务器、Nginx、systemd", "支持通过公网 IP 的 /finance/ 路径访问 Web 系统，后端作为系统服务运行。"),
    ])

    add_heading(doc, "三、角色与权限", 1)
    add_roles_table(doc)
    doc.add_paragraph()
    add_body_paragraph(doc, "系统通过登录状态和角色判断控制菜单显示与接口访问，普通用户无法访问管理员、财务等受限功能。")

    add_heading(doc, "四、核心功能模块", 1)
    add_feature_table(doc, [
        ("登录与角色权限", "登录接口、简单 token、本地登录状态、四类测试账号、角色菜单控制", "体现系统身份识别和基础权限隔离。"),
        ("报销申请管理", "新增、编辑、删除、查询、提交报销单；员工只能操作自己的报销单", "覆盖报销业务的基础入口。"),
        ("审批流程", "草稿、已提交、部门审批中、财务审批中、已通过、已驳回", "完整模拟企业内部多级审批。"),
        ("审批流程时间轴", "详情页展示创建、提交、部门审批、财务审批、最终结果", "老师可以直观看到每张报销单的生命周期。"),
        ("预算管理", "部门预算设置、财务审批通过后自动扣减预算、预算不足提醒", "体现财务控制和业务约束。"),
        ("首页仪表盘", "本月报销金额、待审批数量、报销状态统计、部门预算使用图表", "提升系统展示效果和数据可视化能力。"),
        ("发票附件上传", "上传发票图片或 PDF，附件与报销单关联", "补齐真实报销场景中的凭证管理。"),
        ("OCR 发票识别", "模拟 OCR + 百度 OCR API 接入，识别发票关键字段", "体现第三方 API 集成能力。"),
        ("OCR 金额校验", "比较 OCR 金额与报销金额，显示一致、异常、未识别提醒", "帮助审批人员发现异常，提高财务数据安全性。"),
        ("Excel 导出", "财务和管理员可导出全部报销申请数据为 xlsx", "满足财务报表和离线归档需求。"),
        ("用户管理", "管理员新增、编辑、删除、启用、禁用用户，维护角色和部门", "补充后台系统的管理能力。"),
        ("个人资料", "用户可修改姓名、手机号、邮箱和密码", "提升系统完整度和用户维护能力。"),
        ("操作日志与安全审计", "记录登录、报销、审批、预算、附件、OCR、导出、用户管理等关键操作", "体现“谁在什么时候做了什么”，增强安全审计能力。"),
    ])

    add_heading(doc, "五、数据库设计", 1)
    add_bullets(doc, [
        "user：保存用户账号、密码、角色、部门、联系方式和启用状态。",
        "department：保存部门信息。",
        "reimbursement：保存报销单标题、费用类型、金额、发生日期、状态、申请人、部门、提交时间等。",
        "approval_record：保存审批节点、审批人、审批动作和审批意见。",
        "budget：保存部门年度预算、已用金额和剩余额度。",
        "attachment：保存发票附件文件名、路径、类型、大小和所属报销单。",
        "invoice_ocr_result：保存 OCR 识别字段、确认状态、金额校验结果和原始识别数据。",
        "operation_log：保存关键操作日志，包括用户、角色、模块、操作、对象、详情、IP 和时间。",
    ])

    add_heading(doc, "六、业务流程", 1)
    steps = [
        "员工登录系统，创建报销单并填写费用信息。",
        "员工上传发票附件，执行 OCR 识别并确认识别结果。",
        "系统自动校验发票金额与报销金额是否一致，并在详情页提示。",
        "员工提交报销单，进入部门负责人审批。",
        "部门负责人审批通过后，进入财务审批；如驳回则流程结束。",
        "财务审批通过后，系统检查预算并自动扣减部门预算。",
        "管理员可查看操作日志，追踪关键操作和数据变更。",
    ]
    for index, step in enumerate(steps, start=1):
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(step)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)

    add_heading(doc, "七、项目亮点", 1)
    add_bullets(doc, [
        "功能闭环完整：从报销申请到审批、预算扣减、发票识别、统计展示、日志审计，形成完整业务链路。",
        "权限设计清晰：员工、部门负责人、财务、管理员四类角色职责明确。",
        "有真实业务感：加入预算控制、发票附件、OCR 金额校验、Excel 导出等财务系统常见能力。",
        "可展示性强：首页图表、详情弹窗、审批流程时间轴、红黄绿金额校验提示都适合现场演示。",
        "具备上线能力：项目已支持服务器部署、Nginx 反向代理、systemd 后端服务和公网访问。",
        "体现 Vibe Coding 实践：通过 AI 辅助完成需求拆解、代码生成、报错修复、部署指导和功能迭代。",
    ])

    add_heading(doc, "八、部署与访问", 1)
    add_body_paragraph(doc, "本项目已改造成适合部署到服务器的版本，前端通过 /finance/ 路径访问，后端接口统一走 /api，Nginx 负责静态资源服务和接口反向代理。")
    add_bullets(doc, [
        "本地开发：前端运行 Vite，后端运行 Spring Boot，数据库使用本机 MySQL。",
        "线上部署：后端 jar 放置在 /opt/finance-reimbursement-system/backend，前端 dist 放置在 /var/www/finance-reimbursement-system。",
        "访问方式：http://服务器公网IP/finance/",
        "数据库升级：后端启动后由 JPA 自动补充新增表和字段。",
    ])

    add_heading(doc, "九、演示建议", 1)
    add_bullets(doc, [
        "先用员工账号创建报销单，上传发票并执行 OCR。",
        "修改 OCR 金额制造不一致场景，展示红色金额校验提醒。",
        "提交报销单后切换部门负责人账号审批，展示审批流程时间轴变化。",
        "切换财务账号完成财务审批，展示预算扣减和 Excel 导出。",
        "最后切换管理员账号，展示用户管理和操作日志，说明系统具备安全审计能力。",
    ])

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("财务报销与预算管理系统 | 项目功能汇总")
    footer_run.font.name = "Microsoft YaHei"
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(102, 112, 133)

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
